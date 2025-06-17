import os
import asyncio
import xlrd  # Add this import at the top

from langchain_core.documents import Document
import logging
from concurrent.futures import ProcessPoolExecutor
import functools
import io
from datetime import datetime
import re
import hashlib
import csv

from .. import config
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Import specialized parsers for supported file types
from docx import Document as DocxDocument
import PyPDF2

logger = logging.getLogger(__name__)


def _process_file(file_path: str, chunk_size: int, chunk_overlap: int) -> list[tuple[str, dict]]:
    """
    Processes a single file using specialized parsers for each supported file type.
    Extracts text and simplified metadata for embedding and LLM context.
    """
    logger.info(f"PID {os.getpid()}: Processing file: {file_path}")
    docs_data = []
    
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract text and metadata based on file type
        if file_extension == '.txt':
            text, metadata = _parse_txt(file_path)
        elif file_extension == '.pdf':
            text, metadata = _parse_pdf(file_path)
        elif file_extension == '.docx':
            text, metadata = _parse_docx(file_path)
        elif file_extension == '.csv':
            text, metadata = _parse_csv(file_path)
        elif file_extension == '.xls':
            text, metadata = _parse_xls(file_path)  # Added support for XLS files
        elif file_extension == '.doc':
            text, metadata = _parse_doc(file_path) 
        else:
            logger.warning(f"PID {os.getpid()}: Unsupported file type: {file_extension} for file {file_path}")
            return []
        
        # Create text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        # Split the text if it exceeds chunk size
        if len(text) > chunk_size:
            logger.debug(f"PID {os.getpid()}: Text length ({len(text)}) from {file_path} exceeds chunk size, splitting.")
            chunks = text_splitter.split_text(text)
            for i, chunk_text in enumerate(chunks):
                # Prepare embedding metadata with essential info only
                metadata = _create_embedding_metadata(
                    chunk_text=chunk_text,
                    chunk_index=i,
                    total_chunks=len(chunks),
                    file_path=file_path,
                    content_type=metadata.get("content_type", "text/plain"),
                    document_metadata=metadata
                )
                
                docs_data.append((chunk_text, metadata))
            logger.debug(f"PID {os.getpid()}: Split text into {len(chunks)} chunks.")
        else:
            # Single chunk metadata
            metadata = _create_embedding_metadata(
                chunk_text=text,
                chunk_index=0,
                total_chunks=1,
                file_path=file_path,
                content_type=metadata.get("content_type", "text/plain"),
                document_metadata=metadata
            )
            
            docs_data.append((text, metadata))
        
        logger.info(f"PID {os.getpid()}: Successfully processed {file_path} into {len(docs_data)} document chunks.")
        
    except FileNotFoundError:
        logger.error(f"PID {os.getpid()}: File not found: {file_path}")
    except Exception as e:
        logger.error(f"PID {os.getpid()}: Error processing file {file_path}: {e}", exc_info=True)
    
    return docs_data


def _create_embedding_metadata(chunk_text: str, chunk_index: int, total_chunks: int, 
                              file_path: str, content_type: str, document_metadata: dict) -> dict:
    """
    Creates simplified metadata for embedding and LLM context retrieval.
    Now includes the raw content as part of the metadata.
    """
    # Basic info - just filename and position
    filename = os.path.basename(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Simplified essential metadata, now including the raw content
    embedding_metadata = {
        "filename": filename,
        "chunk_index": chunk_index,
        "total_chunks": total_chunks,
        "raw_content": chunk_text  # Include the raw content in metadata
    }
    
    # Add page information for PDFs (if available)
    if file_extension == '.pdf' and "page_boundaries" in document_metadata:
        # Find which page this chunk belongs to based on text position
        chunk_start_pos = 0
        for prev_chunk_idx in range(chunk_index):
            prev_chunk_len = len(document_metadata.get("chunk_texts", [""])[prev_chunk_idx])
            chunk_start_pos += prev_chunk_len
        
        page_boundaries = document_metadata["page_boundaries"]
        for page_num, boundary in enumerate(page_boundaries):
            if chunk_start_pos <= boundary:
                embedding_metadata["page"] = page_num + 1
                break
    
    # Add document metadata that's genuinely useful
    if document_metadata.get("author"):
        embedding_metadata["author"] = document_metadata.get("author")
    
    if document_metadata.get("title"):
        embedding_metadata["title"] = document_metadata.get("title")
    
    # Extract any dates found in the text (usually important context)
    date_patterns = [
        r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
        r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY/MM/DD
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'  # Month DD, YYYY
    ]
    
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, chunk_text)
        dates.extend(matches)
    
    if dates:
        embedding_metadata["dates"] = dates[:3]  # Only include up to 3 dates
    
    return embedding_metadata


def _extract_basic_entities(text: str) -> dict:
    """Extract basic entities from text for better context retrieval."""
    entities = {
        "dates": [],
        "numbers": [],
        "emails": [],
        "urls": []
    }
    
    # Simple pattern matching for dates (various formats)
    date_patterns = [
        r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
        r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY/MM/DD
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'  # Month DD, YYYY
    ]
    
    for pattern in date_patterns:
        matches = re.findall(pattern, text)
        entities["dates"].extend(matches)
    
    # Extract numbers (including those with decimal points)
    number_pattern = r'\b\d+(?:\.\d+)?\b'
    entities["numbers"] = re.findall(number_pattern, text)[:10]  # Limit to first 10
    
    # Extract emails
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    entities["emails"] = re.findall(email_pattern, text)
    
    # Extract URLs
    url_pattern = r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+'
    entities["urls"] = re.findall(url_pattern, text)
    
    return entities


def _detect_semantic_structure(text: str) -> str:
    """Detect the semantic structure of the text chunk."""
    # Check for headings (lines ending with : or starting with # or numbers)
    heading_patterns = [
        r'^#+\s+.+$',                # Markdown heading
        r'^[A-Z][A-Za-z\s]+:',       # Title followed by colon
        r'^\d+\.\s+[A-Z]',           # Numbered item
        r'^[IVXLCDM]+\.\s+[A-Z]'     # Roman numeral heading
    ]
    
    for line in text.split('\n'):
        for pattern in heading_patterns:
            if re.match(pattern, line.strip()):
                return "heading_with_content"
    
    # Check for bullet points
    if re.search(r'^\s*[\*\-\•]\s+', text, re.MULTILINE):
        return "bullet_list"
    
    # Check for numbered lists
    if re.search(r'^\s*\d+\.\s+', text, re.MULTILINE):
        return "numbered_list"
    
    # Check for Q&A pattern
    if re.search(r'\b(?:Q|Question)[.:]\s+', text, re.IGNORECASE) and re.search(r'\b(?:A|Answer)[.:]\s+', text, re.IGNORECASE):
        return "question_answer"
    
    # Check for tabular data (multiple | characters or consistent spacing)
    if re.search(r'[|\t]', text) and re.search(r'\n.*[|\t]', text):
        return "tabular_data"
    
    # Default to paragraph if nothing else matches
    return "paragraph"


def _parse_txt(file_path: str) -> tuple[str, dict]:
    """Parse a text file and extract its content."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
        text = file.read()
    
    # Extract basic metadata
    metadata = {
        "content_type": "plain/text",
        "line_count": len(text.splitlines()),
        "char_count": len(text)
    }
    
    return text, metadata


def _parse_pdf(file_path: str) -> tuple[str, dict]:
    """Parse a PDF file and extract its content and metadata."""
    full_text = ""
    metadata = {
        "content_type": "application/pdf",
        "page_count": 0,
        "has_images": False,
        "pdf_version": None
    }
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        metadata["page_count"] = len(pdf_reader.pages)
        metadata["pdf_version"] = pdf_reader.pdf_header
        
        # Check if PDF has images (basic check)
        for page in pdf_reader.pages:
            if '/XObject' in page:
                metadata["has_images"] = True
                break
        
        # Extract text from each page
        page_texts = []
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
            full_text += page_text + "\n\n"
        
        # Store page texts for context retrieval
        metadata["chunk_texts"] = page_texts
        
        # Store page boundaries for better context retrieval
        metadata["page_boundaries"] = [len("".join(page_texts[:i+1])) for i in range(len(page_texts))]
        
        # Try to extract document info if available
        if pdf_reader.metadata:
            if pdf_reader.metadata.get('/Title'):
                metadata["title"] = pdf_reader.metadata.get('/Title')
            if pdf_reader.metadata.get('/Author'):
                metadata["author"] = pdf_reader.metadata.get('/Author')
            if pdf_reader.metadata.get('/Subject'):
                metadata["subject"] = pdf_reader.metadata.get('/Subject')
            if pdf_reader.metadata.get('/Keywords'):
                metadata["keywords"] = pdf_reader.metadata.get('/Keywords')
    
    return full_text, metadata

def _parse_doc(file_path: str) -> tuple[str, dict]:
    """Parse a DOC file and extract its content and metadata."""
    try:
        from docx import Document as DocDocument
    except ImportError:
        raise ImportError("Please install python-docx to parse DOC files: pip install python-docx")
    
    doc = DocDocument(file_path)
    
    # Extract text from paragraphs
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    
    # Extract metadata
    metadata = {
        "content_type": "application/msword",
        "paragraph_count": len(paragraphs),
        "has_tables": len(doc.tables) > 0,
        "table_count": len(doc.tables)
    }
    
    # Extract text from tables if present
    if metadata["has_tables"]:
        table_texts = []
        for table in doc.tables:
            table_text = ""
            for i, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                table_text += row_text + "\n"
            table_texts.append(table_text)
        
        # Add table text to full text
        tables_content = "\n\n".join(table_texts)
        full_text += "\n\n" + tables_content
        
        # Store table data for better context
        metadata["table_texts"] = table_texts
    
    # Extract document properties if available
    try:
        core_props = doc.core_properties
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created_date"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modified_date"] = core_props.modified.isoformat()
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.subject:
            metadata["subject"] = core_props.subject
    except Exception as e:
        logger.error(f"Error extracting DOC properties: {e}", exc_info=True)
    
    return full_text, metadata

def _parse_docx(file_path: str) -> tuple[str, dict]:
    """Parse a DOCX file and extract its content and metadata."""
    doc = DocxDocument(file_path)
    
    # Extract text from paragraphs
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    
    # Extract metadata
    metadata = {
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "paragraph_count": len(paragraphs),
        "has_tables": len(doc.tables) > 0,
        "table_count": len(doc.tables)
    }
    
    # Extract text from tables if present
    if metadata["has_tables"]:
        table_texts = []
        for table in doc.tables:
            table_text = ""
            for i, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                table_text += row_text + "\n"
            table_texts.append(table_text)
        
        # Add table text to full text
        tables_content = "\n\n".join(table_texts)
        full_text += "\n\n" + tables_content
        
        # Store table data for better context
        metadata["table_texts"] = table_texts
    
    # Extract document properties if available
    try:
        core_props = doc.core_properties
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created_date"] = core_props.created.isoformat()
        if core_props.modified:
            metadata["modified_date"] = core_props.modified.isoformat()
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.keywords:
            metadata["keywords"] = core_props.keywords
        if core_props.subject:
            metadata["subject"] = core_props.subject
    except:
        pass
    
    # Extract headings (basic implementation)
    headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            headings.append(paragraph.text)
    
    if headings:
        metadata["headings"] = headings
    
    return full_text, metadata


def _parse_csv(file_path: str) -> tuple[str, dict]:
    """Parse a CSV file and extract its content and metadata using built-in csv module."""
    
    # First try to detect the delimiter using csv.Sniffer
    delimiter = ','
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            sample = f.read(4096)
            if sample:
                sniffer = csv.Sniffer()
                dialect = sniffer.sniff(sample)
                delimiter = dialect.delimiter
                has_header = sniffer.has_header(sample)
    except Exception as e:
        logger.warning(f"Could not detect CSV delimiter, using comma as default: {e}")
        has_header = True  # Assume header by default
    
    # Read the CSV file with the built-in csv module
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            csv_reader = csv.reader(file, delimiter=delimiter)
            rows = list(csv_reader)
            
            if not rows:
                return "", {"content_type": "text/csv", "empty": True}
            
            headers = rows[0] if rows and has_header else []
            data_rows = rows[1:] if rows and has_header else rows
            
            # Format as text with proper tabular representation for better readability
            text_buffer = io.StringIO()
            
            # Create a formatted tabular representation
            tabular_text = ""
            
            # If we have headers, use them
            if headers and has_header:
                tabular_text += " | ".join(str(h) for h in headers) + "\n"
                tabular_text += "-" * (sum(len(str(h)) + 3 for h in headers)) + "\n"
                
                # Format the main text with headers
                text_buffer.write(",".join(headers) + "\n")
            
            # Add rows (limit to a reasonable number to avoid excessive text)
            max_rows_to_include = min(100, len(data_rows))
            for i, row in enumerate(data_rows):
                if i < max_rows_to_include:
                    tabular_text += " | ".join(str(v) for v in row) + "\n"
                    text_buffer.write(",".join(row) + "\n")
            
            if len(data_rows) > max_rows_to_include:
                tabular_text += f"\n... and {len(data_rows) - max_rows_to_include} more rows ...\n"
            
            text = text_buffer.getvalue()
            
            # Detect column types by examining the data
            column_types = {}
            if headers and data_rows:
                for col_idx, col_name in enumerate(headers):
                    # Get a sample of values for this column
                    sample_values = []
                    for row_idx, row in enumerate(data_rows):
                        if row_idx < 10 and col_idx < len(row):  # Sample from first 10 rows
                            sample_values.append(row[col_idx])
                    
                    # Determine column type
                    col_type = "string"  # Default type
                    
                    # Check if all values are numeric
                    is_numeric = True
                    has_decimal = False
                    
                    for val in sample_values:
                        if not val:  # Skip empty values
                            continue
                        # Try to convert to float
                        try:
                            num = float(val)
                            if '.' in val:
                                has_decimal = True
                        except ValueError:
                            is_numeric = False
                            break
                    
                    if is_numeric:
                        col_type = "float" if has_decimal else "integer"
                        
                    # Check if looks like dates (simple check)
                    date_patterns = [
                        r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
                        r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY/MM/DD
                    ]
                    
                    looks_like_date = False
                    for val in sample_values:
                        if not val:
                            continue
                        for pattern in date_patterns:
                            if re.match(pattern, val):
                                looks_like_date = True
                                break
                        if looks_like_date:
                            break
                    
                    if looks_like_date:
                        col_type = "datetime"
                    
                    column_types[col_name] = col_type
            
            # Basic metadata extraction
            metadata = {
                "content_type": "text/csv",
                "column_count": len(headers) if has_header else len(rows[0]) if rows else 0,
                "row_count": len(data_rows),
                "columns": headers if has_header else [],
                "has_header": has_header,
                "delimiter": delimiter,
                "tabular_representation": tabular_text
            }
            
            # Add column types if available
            if column_types:
                metadata["column_types"] = column_types
            
            # Add sample rows
            sample_rows = []
            max_samples = min(5, len(data_rows))
            
            if has_header and headers:
                for i in range(max_samples):
                    if i < len(data_rows):
                        row_dict = {}
                        for j, header in enumerate(headers):
                            if j < len(data_rows[i]):
                                row_dict[header] = data_rows[i][j]
                        sample_rows.append(row_dict)
                metadata["sample_rows"] = sample_rows
            
            # Look for potential date columns based on naming
            if has_header:
                date_columns = [col for col in headers if any(date_term in col.lower() 
                                                         for date_term in ['date', 'time', 'year', 'month', 'day'])]
                if date_columns:
                    metadata["potential_date_columns"] = date_columns
            
            # Try to extract basic statistics for numeric columns
            if has_header and column_types:
                numeric_stats = {}
                for col_idx, col_name in enumerate(headers):
                    if column_types.get(col_name) in ["integer", "float"]:
                        # Get all values for this column
                        values = []
                        null_count = 0
                        for row in data_rows:
                            if col_idx < len(row):
                                val = row[col_idx]
                                if val.strip():  # Skip empty values
                                    try:
                                        values.append(float(val))
                                    except ValueError:
                                        pass
                                else:
                                    null_count += 1
                        
                        if values:
                            # Calculate simple statistics
                            col_stats = {
                                "min": min(values),
                                "max": max(values),
                                "mean": sum(values) / len(values),
                                "null_count": null_count
                            }
                            # Sort values to find median
                            values.sort()
                            if len(values) % 2 == 0:
                                median = (values[len(values)//2 - 1] + values[len(values)//2]) / 2
                            else:
                                median = values[len(values)//2]
                            col_stats["median"] = median
                            
                            numeric_stats[col_name] = col_stats
                
                if numeric_stats:
                    metadata["numeric_stats"] = numeric_stats
            
            return text, metadata
    
    except Exception as e:
        logger.error(f"Error parsing CSV file {file_path}: {e}")
        # Return a minimal representation in case of errors
        return file_path, {"content_type": "text/csv", "error": str(e)}


def _parse_xls(file_path: str) -> tuple[str, dict]:
    """Parse an XLS file and extract its content and metadata."""
    try:
        workbook = xlrd.open_workbook(file_path)
        sheet = workbook.sheet_by_index(0)  # Read the first sheet
        
        # Extract data
        rows = []
        for row_idx in range(sheet.nrows):
            row = [sheet.cell_value(row_idx, col_idx) for col_idx in range(sheet.ncols)]
            rows.append(row)
        
        # Format as text
        text = "\n".join(["\t".join(map(str, row)) for row in rows])
        
        # Metadata
        metadata = {
            "content_type": "application/vnd.ms-excel",
            "sheet_count": len(workbook.sheets()),
            "row_count": sheet.nrows,
            "column_count": sheet.ncols,
        }
        return text, metadata
    except Exception as e:
        logger.error(f"Error parsing XLS file {file_path}: {e}")
        return "", {"content_type": "application/vnd.ms-excel", "error": str(e)}


class AsyncFileParser:
    def __init__(self, supported_extensions=None):
        self.supported_extensions = supported_extensions or config.SUPPORTED_EXTENSIONS 
        logger.info(f"AsyncFileParser initialized. Will use specialized parsers for: {self.supported_extensions}")

        self.process_pool_executor = None
        if not config.DISABLE_FILE_PROCESSING_PARALLELISM:
            logger.info(f"Initializing ProcessPoolExecutor for parallel file processing (default workers: {os.cpu_count()}).")
            self.process_pool_executor = ProcessPoolExecutor() # Uses os.cpu_count() by default
        else:
            logger.info("Parallel file processing is DISABLED. CPU-bound tasks will run in the main process.")

    async def _run_cpu_bound_processing(self, file_path: str):
        if self.process_pool_executor:
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(
                self.process_pool_executor, 
                functools.partial(_process_file, file_path, config.CHUNK_SIZE, config.CHUNK_OVERLAP)
            )
        else:
            # Parallelism is disabled, run directly (will block the event loop for this task)
            logger.debug(f"Running _process_file for {file_path} sequentially in main process.")
            try:
                return _process_file(file_path, config.CHUNK_SIZE, config.CHUNK_OVERLAP)
            except Exception as e:
                logger.error(f"Error during sequential execution of _process_file for {file_path}: {e}", exc_info=True)
                return [] # Return empty list on error, consistent with executor path

    async def load_and_process_file(self, file_path: str) -> list[Document]:
        logger.info(f"Queueing file for processing: {file_path}")
        try:
            # _process_file now returns list of (text, metadata_dict) tuples
            processed_data_list = await self._run_cpu_bound_processing(file_path)
            
            if processed_data_list is None: # Should generally not happen if worker returns list, even empty
                logger.error(f"Worker for {file_path} returned None. Investigate _process_file behavior.")
                return []

            documents = [Document(page_content=text, metadata=meta) for text, meta in processed_data_list]
            logger.info(f"Successfully created {len(documents)} Langchain Documents from {file_path}.")
            return documents
        except Exception as e_load_process:
            logger.error(f"Error in load_and_process_file for {file_path}: {e_load_process}", exc_info=True)
            return []

    async def process_path_intelligently(self, input_path: str) -> list[Document]:
        all_documents = []
        abs_input_path = os.path.abspath(input_path)
        if not os.path.exists(abs_input_path):
            logger.error(f"Input path does not exist: {abs_input_path}")
            return []

        files_to_process = []
        if os.path.isfile(abs_input_path):
            logger.info(f"Input is a single file: {abs_input_path}. Submitting for processing.")
            file_ext = os.path.splitext(abs_input_path)[1].lower()
            if file_ext in self.supported_extensions:
                files_to_process.append(abs_input_path)
            else:
                logger.warning(f"File {abs_input_path} has unsupported extension {file_ext}. Skipping.")

        elif os.path.isdir(abs_input_path):
            logger.info(f"Input is a directory: {abs_input_path}. Scanning for files...")
            for root, _, files in os.walk(abs_input_path):
                for file_name in files:
                    file_path = os.path.join(root, file_name)
                    file_ext = os.path.splitext(file_path)[1].lower()
                    if file_ext in self.supported_extensions:
                        files_to_process.append(os.path.abspath(file_path))
                    else:
                        logger.debug(f"Skipping file with unsupported extension: {file_path}")
        else:
            logger.warning(f"Input path {abs_input_path} is neither a file nor a directory.")
            return []

        if not files_to_process:
            logger.info(f"No files found to process for input: {abs_input_path} (after filtering for supported extensions).")
            return []

        logger.info(f"Found {len(files_to_process)} files. Submitting to process via asyncio.gather.")
        tasks = [self.load_and_process_file(fp) for fp in files_to_process]
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        successful_files_count = 0
        failed_files_count = 0
        for i, result_item in enumerate(results):
            current_file_path = files_to_process[i]
            if isinstance(result_item, Exception):
                logger.error(f"Error processing file {current_file_path} during gather: {result_item}", exc_info=True)
                failed_files_count += 1
            elif isinstance(result_item, list):
                all_documents.extend(result_item)
                successful_files_count += 1
                if not result_item:
                     logger.info(f"File {current_file_path} processed successfully but yielded 0 documents (e.g., empty file).")
            else:
                logger.error(f"Unexpected result type for file {current_file_path}: {type(result_item)}.")
                failed_files_count += 1
        
        logger.info(
            f"Processing batch from {abs_input_path} finished. "
            f"Successfully processed files: {successful_files_count}. "
            f"Failed files: {failed_files_count}. "
            f"Total Langchain Documents generated: {len(all_documents)}."
        )
        return all_documents

    def close(self):
        if hasattr(self, 'process_pool_executor') and self.process_pool_executor:
            logger.info("Shutting down ProcessPoolExecutor in AsyncFileParser...")
            try:
                self.process_pool_executor.shutdown(wait=True, cancel_futures=False)
                logger.info("ProcessPoolExecutor shut down successfully.")
            except Exception as e_shutdown:
                logger.error(f"Error during ProcessPoolExecutor shutdown: {e_shutdown}", exc_info=True)
            self.process_pool_executor = None
        else:
            logger.info("ProcessPoolExecutor was not initialized or already shut down.")